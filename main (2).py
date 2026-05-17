"""
╔══════════════════════════════════════════════════════════════════════╗
║         ISHXONA ISH VAQTI BOTI — v2.0                              ║
║         python-telegram-bot==21.10 | Python 3.13+                  ║
║         Bot: @ishxona_ishvaqtibot                                   ║
║         Developer ID: 8346931722                                    ║
╚══════════════════════════════════════════════════════════════════════╝
"""

import asyncio
import csv
import io
import logging
import os
import sqlite3
from datetime import datetime, timedelta
from typing import Optional

import pytz
from docx import Document
from docx.shared import Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

from telegram import (
    Update,
    InlineKeyboardButton,
    InlineKeyboardMarkup,
)
from telegram.constants import ParseMode
from telegram.ext import (
    Application,
    CallbackQueryHandler,
    CommandHandler,
    ContextTypes,
    MessageHandler,
    filters,
)

# ═══════════════════════════════════════════════════════════════
#                        KONFIGURATSIYA
# ═══════════════════════════════════════════════════════════════

BOT_TOKEN        = os.environ.get("BOT_TOKEN", "8814493938:AAHzLkvvhS3nRM61bkPSOD59qd6WaGHGBbk")
SUPER_ADMIN_ID   = int(os.environ.get("SUPER_ADMIN_ID",   "8346931722"))
DEFAULT_ADMIN_ID = int(os.environ.get("DEFAULT_ADMIN_ID", "7918082766"))

DB_PATH      = "attendance.db"
TASHKENT_TZ  = pytz.timezone("Asia/Tashkent")

WORK_START_H, WORK_START_M = 8, 30
WORK_END_H,   WORK_END_M   = 20, 0

STATE          = "state"
WARN_WORKER_ID = "warn_worker_id"
S_ADD_WORKER   = "add_worker"
S_ADD_ADMIN    = "add_admin"
S_ADD_MANAGER  = "add_manager"
S_BROADCAST    = "broadcast"
S_WARN_MSG     = "warn_msg"

logging.basicConfig(
    format="%(asctime)s | %(levelname)s | %(message)s",
    level=logging.INFO,
    handlers=[logging.StreamHandler(), logging.FileHandler("bot.log", encoding="utf-8")],
)
log = logging.getLogger(__name__)


# ═══════════════════════════════════════════════════════════════
#                     MA'LUMOTLAR BAZASI
# ═══════════════════════════════════════════════════════════════

def db() -> sqlite3.Connection:
    conn = sqlite3.connect(DB_PATH, check_same_thread=False)
    conn.row_factory = sqlite3.Row
    return conn


def init_db() -> None:
    with db() as conn:
        conn.executescript("""
            CREATE TABLE IF NOT EXISTS workers (
                id          INTEGER PRIMARY KEY AUTOINCREMENT,
                name        TEXT    NOT NULL,
                telegram_id INTEGER UNIQUE,
                username    TEXT,
                added_date  TEXT    NOT NULL,
                is_active   INTEGER DEFAULT 1
            );
            CREATE TABLE IF NOT EXISTS attendance (
                id           INTEGER PRIMARY KEY AUTOINCREMENT,
                worker_id    INTEGER NOT NULL,
                action       TEXT    NOT NULL CHECK(action IN ('in','out')),
                time         TEXT    NOT NULL,
                date         TEXT    NOT NULL,
                is_late      INTEGER DEFAULT 0,
                late_minutes INTEGER DEFAULT 0,
                FOREIGN KEY (worker_id) REFERENCES workers(id)
            );
            CREATE TABLE IF NOT EXISTS admins (
                id          INTEGER PRIMARY KEY AUTOINCREMENT,
                telegram_id INTEGER UNIQUE NOT NULL,
                name        TEXT,
                role        TEXT NOT NULL CHECK(role IN ('superadmin','admin','manager')),
                added_date  TEXT NOT NULL
            );
            CREATE TABLE IF NOT EXISTS warnings (
                id          INTEGER PRIMARY KEY AUTOINCREMENT,
                worker_id   INTEGER NOT NULL,
                message     TEXT    NOT NULL,
                sent_date   TEXT    NOT NULL,
                sent_by     INTEGER,
                FOREIGN KEY (worker_id) REFERENCES workers(id)
            );
        """)
        now = now_tashkent().strftime("%Y-%m-%d %H:%M:%S")
        conn.execute(
            "INSERT OR IGNORE INTO admins(telegram_id,name,role,added_date) VALUES(?,?,?,?)",
            (SUPER_ADMIN_ID, "Developer", "superadmin", now),
        )
        conn.execute(
            "INSERT OR IGNORE INTO admins(telegram_id,name,role,added_date) VALUES(?,?,?,?)",
            (DEFAULT_ADMIN_ID, "Qahramon Qalandarov", "admin", now),
        )
    log.info("Ma'lumotlar bazasi tayyor.")


# ═══════════════════════════════════════════════════════════════
#                     YORDAMCHI FUNKSIYALAR
# ═══════════════════════════════════════════════════════════════

def now_tashkent() -> datetime:
    return datetime.now(TASHKENT_TZ)

def today_str() -> str:
    return now_tashkent().strftime("%Y-%m-%d")

def fmt_dt(dt_str: str) -> str:
    try:
        return datetime.strptime(dt_str, "%Y-%m-%d %H:%M:%S").strftime("%d.%m.%Y %H:%M")
    except Exception:
        return dt_str

MONTHS = ["Yanvar","Fevral","Mart","Aprel","May","Iyun",
          "Iyul","Avgust","Sentabr","Oktabr","Noyabr","Dekabr"]
DAYS   = ["Dushanba","Seshanba","Chorshanba","Payshanba","Juma","Shanba","Yakshanba"]

def fmt_date_full(date_str: str) -> str:
    try:
        d = datetime.strptime(date_str, "%Y-%m-%d")
        return f"{d.day} {MONTHS[d.month-1]} {d.year} ({DAYS[d.weekday()]})"
    except Exception:
        return date_str

def hm(minutes: int) -> str:
    h, m = divmod(minutes, 60)
    return f"{h} soat {m} daqiqa" if h else f"{m} daqiqa"

def late_minutes_calc(time_str: str) -> int:
    try:
        t     = datetime.strptime(time_str, "%Y-%m-%d %H:%M:%S")
        start = t.replace(hour=WORK_START_H, minute=WORK_START_M, second=0, microsecond=0)
        if t > start:
            return int((t - start).total_seconds() / 60)
    except Exception:
        pass
    return 0

def get_role(uid: int) -> Optional[str]:
    with db() as conn:
        row = conn.execute("SELECT role FROM admins WHERE telegram_id=?", (uid,)).fetchone()
    return row["role"] if row else None

def is_admin(uid: int) -> bool:
    return get_role(uid) is not None

def all_admin_ids() -> list:
    with db() as conn:
        rows = conn.execute("SELECT telegram_id FROM admins").fetchall()
    return [r["telegram_id"] for r in rows]


# ═══════════════════════════════════════════════════════════════
#                         KLAVIATURALAR
# ═══════════════════════════════════════════════════════════════

def kb_main(uid: int) -> InlineKeyboardMarkup:
    rows = [
        [InlineKeyboardButton("⏱ Keldi / Ketdi", callback_data="att_menu")],
        [InlineKeyboardButton("📊 Hisobot",       callback_data="rep_menu")],
    ]
    if is_admin(uid):
        rows.append([InlineKeyboardButton("👥 Admin Panel", callback_data="adm_panel")])
    return InlineKeyboardMarkup(rows)

def kb_report() -> InlineKeyboardMarkup:
    return InlineKeyboardMarkup([
        [InlineKeyboardButton("📅 Bugungi",  callback_data="rep_today")],
        [InlineKeyboardButton("📆 Haftalik", callback_data="rep_week")],
        [InlineKeyboardButton("🗓 Oylik",    callback_data="rep_month")],
        [InlineKeyboardButton("📎 CSV",  callback_data="exp_csv"),
         InlineKeyboardButton("📄 Word", callback_data="exp_word")],
        [InlineKeyboardButton("🏠 Asosiy", callback_data="main_menu")],
    ])

def kb_admin(uid: int) -> InlineKeyboardMarkup:
    role = get_role(uid)
    rows = [
        [InlineKeyboardButton("➕ Ishchi qo'shish",  callback_data="add_worker"),
         InlineKeyboardButton("🗑 Ishchi o'chirish", callback_data="del_worker_list")],
        [InlineKeyboardButton("👁 Barcha ishchilar", callback_data="all_workers")],
        [InlineKeyboardButton("📢 Xabar (broadcast)", callback_data="broadcast")],
        [InlineKeyboardButton("⚠️ Ogohlantirish",   callback_data="warn_list")],
    ]
    if role in ("superadmin", "admin"):
        rows += [
            [InlineKeyboardButton("👤 Admin qo'shish",   callback_data="add_admin"),
             InlineKeyboardButton("🗑 Admin o'chirish",  callback_data="del_admin_list")],
            [InlineKeyboardButton("👨 Manager qo'shish", callback_data="add_manager"),
             InlineKeyboardButton("🗑 Manager o'chirish",callback_data="del_manager_list")],
            [InlineKeyboardButton("📋 Adminlar ro'yxati",callback_data="list_admins")],
        ]
    rows.append([InlineKeyboardButton("🏠 Asosiy menyu", callback_data="main_menu")])
    return InlineKeyboardMarkup(rows)

def kb_back(cb: str = "main_menu") -> InlineKeyboardMarkup:
    return InlineKeyboardMarkup([[InlineKeyboardButton("⬅️ Orqaga", callback_data=cb)]])


# ═══════════════════════════════════════════════════════════════
#                          /start
# ═══════════════════════════════════════════════════════════════

async def cmd_start(update: Update, ctx: ContextTypes.DEFAULT_TYPE) -> None:
    user = update.effective_user
    uid  = user.id
    now  = now_tashkent()

    with db() as conn:
        conn.execute(
            "UPDATE workers SET telegram_id=?, username=? WHERE telegram_id=?",
            (uid, user.username, uid),
        )

    hour  = now.hour
    greet = ("🌅 Xayrli tong" if 5 <= hour < 12 else
             "☀️ Xayrli kun"  if 12 <= hour < 17 else
             "🌆 Xayrli kech" if 17 <= hour < 21 else
             "🌙 Xayrli tun")

    role_map = {"superadmin": "👑 Developer", "admin": "🔑 Admin", "manager": "👨‍💼 Manager"}
    role_txt = role_map.get(get_role(uid), "👤 Xodim")

    text = (
        f"{greet}, <b>{user.first_name}</b>!\n\n"
        f"🏢 <b>Ishxona Ish Vaqti Tizimi</b>\n"
        f"{'━'*22}\n"
        f"📌 Rol: <b>{role_txt}</b>\n"
        f"🕐 Vaqt: <b>{now.strftime('%H:%M')}</b>\n"
        f"📅 Sana: <b>{fmt_date_full(now.strftime('%Y-%m-%d'))}</b>\n"
        f"{'━'*22}\n"
        f"⏰ Ish vaqti: <b>{WORK_START_H:02d}:{WORK_START_M:02d}"
        f" – {WORK_END_H:02d}:{WORK_END_M:02d}</b>\n\n"
        f"Kerakli bo'limni tanlang:"
    )
    await update.message.reply_text(text, reply_markup=kb_main(uid), parse_mode=ParseMode.HTML)


# ═══════════════════════════════════════════════════════════════
#                   ASOSIY CALLBACK DISPATCHER
# ═══════════════════════════════════════════════════════════════

async def on_callback(update: Update, ctx: ContextTypes.DEFAULT_TYPE) -> None:
    q   = update.callback_query
    uid = q.from_user.id
    d   = q.data
    await q.answer()

    try:
        if d == "main_menu":
            ctx.user_data.clear()
            now = now_tashkent()
            await q.edit_message_text(
                f"🏢 <b>Ishxona Ish Vaqti Tizimi</b>\n"
                f"🕐 {now.strftime('%H:%M')}  📅 {fmt_date_full(now.strftime('%Y-%m-%d'))}\n\n"
                f"Bo'limni tanlang:",
                reply_markup=kb_main(uid), parse_mode=ParseMode.HTML,
            )

        elif d == "att_menu":
            await show_att_menu(q)

        elif d.startswith("ws_"):
            await show_worker_actions(q, int(d[3:]))

        elif d.startswith("in_"):
            await record(q, ctx, int(d[3:]), "in")

        elif d.startswith("out_"):
            await record(q, ctx, int(d[4:]), "out")

        elif d == "rep_menu":
            if not is_admin(uid):
                await q.edit_message_text("❌ Ruxsat yo'q.", reply_markup=kb_back())
                return
            await q.edit_message_text(
                "📊 <b>Hisobotlar</b>\n\nQaysi hisobotni ko'rmoqchisiz?",
                reply_markup=kb_report(), parse_mode=ParseMode.HTML,
            )

        elif d in ("rep_today", "rep_week", "rep_month"):
            if not is_admin(uid): return
            await show_report(q, d.split("_")[1])

        elif d == "exp_csv":
            if not is_admin(uid): return
            await send_csv(q, ctx)

        elif d == "exp_word":
            if not is_admin(uid): return
            await send_word(q, ctx)

        elif d == "adm_panel":
            if not is_admin(uid):
                await q.edit_message_text("❌ Admin huquqi yo'q.")
                return
            role_map = {"superadmin": "👑 Developer", "admin": "🔑 Admin", "manager": "👨‍💼 Manager"}
            await q.edit_message_text(
                f"👥 <b>Admin Panel</b>\n"
                f"Rolingiz: <b>{role_map.get(get_role(uid),'')}</b>\n\nAmalni tanlang:",
                reply_markup=kb_admin(uid), parse_mode=ParseMode.HTML,
            )

        elif d == "all_workers":
            if not is_admin(uid): return
            await show_all_workers(q)

        elif d == "list_admins":
            if not is_admin(uid): return
            await show_all_admins(q)

        elif d == "add_worker":
            if not is_admin(uid): return
            ctx.user_data[STATE] = S_ADD_WORKER
            await q.edit_message_text(
                "➕ <b>Yangi ishchi qo'shish</b>\n\nIshchining <b>Ism Familiyasini</b> kiriting:",
                reply_markup=kb_back("adm_panel"), parse_mode=ParseMode.HTML,
            )

        elif d == "add_admin":
            if get_role(uid) not in ("superadmin", "admin"): return
            ctx.user_data[STATE] = S_ADD_ADMIN
            await q.edit_message_text(
                "👤 <b>Admin qo'shish</b>\n\nAdmin Telegram ID sini kiriting:",
                reply_markup=kb_back("adm_panel"), parse_mode=ParseMode.HTML,
            )

        elif d == "add_manager":
            if get_role(uid) not in ("superadmin", "admin"): return
            ctx.user_data[STATE] = S_ADD_MANAGER
            await q.edit_message_text(
                "👨 <b>Manager qo'shish</b>\n\nManager Telegram ID sini kiriting:",
                reply_markup=kb_back("adm_panel"), parse_mode=ParseMode.HTML,
            )

        elif d == "del_worker_list":
            if not is_admin(uid): return
            await show_del_workers(q)

        elif d.startswith("dw_"):
            if not is_admin(uid): return
            await delete_worker(q, int(d[3:]))

        elif d in ("del_admin_list", "del_manager_list"):
            if get_role(uid) not in ("superadmin", "admin"): return
            await show_del_admins(q, "admin" if d == "del_admin_list" else "manager")

        elif d.startswith("da_"):
            if get_role(uid) not in ("superadmin", "admin"): return
            await delete_admin(q, int(d[3:]))

        elif d == "broadcast":
            if not is_admin(uid): return
            ctx.user_data[STATE] = S_BROADCAST
            await q.edit_message_text(
                "📢 <b>Hammaga xabar</b>\n\nXabar matnini yozing:",
                reply_markup=kb_back("adm_panel"), parse_mode=ParseMode.HTML,
            )

        elif d == "warn_list":
            if not is_admin(uid): return
            await show_warn_list(q)

        elif d.startswith("ww_"):
            if not is_admin(uid): return
            ctx.user_data[WARN_WORKER_ID] = int(d[3:])
            ctx.user_data[STATE] = S_WARN_MSG
            await q.edit_message_text(
                "⚠️ <b>Ogohlantirish</b>\n\nOgohlantirish matnini yozing:",
                reply_markup=kb_back("adm_panel"), parse_mode=ParseMode.HTML,
            )

    except Exception as e:
        log.exception(f"Callback xatosi [{d}]: {e}")
        try:
            await q.edit_message_text(
                f"❌ Xatolik: <code>{str(e)[:200]}</code>",
                reply_markup=kb_back(), parse_mode=ParseMode.HTML,
            )
        except Exception:
            pass


# ═══════════════════════════════════════════════════════════════
#                   KELDI / KETDI
# ═══════════════════════════════════════════════════════════════

async def show_att_menu(q) -> None:
    with db() as conn:
        workers = conn.execute(
            "SELECT id, name FROM workers WHERE is_active=1 ORDER BY name"
        ).fetchall()

    if not workers:
        await q.edit_message_text(
            "👥 Ishchilar yo'q.\nAdmin panel orqali qo'shing.",
            reply_markup=kb_back(),
        )
        return

    btns = [[InlineKeyboardButton(f"👤 {w['name']}", callback_data=f"ws_{w['id']}")]
             for w in workers]
    btns.append([InlineKeyboardButton("🏠 Asosiy", callback_data="main_menu")])
    await q.edit_message_text(
        "⏱ <b>Keldi / Ketdi</b>\n\nIsmingizni tanlang:",
        reply_markup=InlineKeyboardMarkup(btns), parse_mode=ParseMode.HTML,
    )


async def show_worker_actions(q, worker_id: int) -> None:
    with db() as conn:
        w    = conn.execute("SELECT name FROM workers WHERE id=?", (worker_id,)).fetchone()
        last = conn.execute(
            "SELECT action, time FROM attendance WHERE worker_id=? AND date=? ORDER BY id DESC LIMIT 1",
            (worker_id, today_str()),
        ).fetchone()

    if not w:
        await q.edit_message_text("❌ Ishchi topilmadi.", reply_markup=kb_back("att_menu"))
        return

    status = ""
    if last:
        a_txt  = "✅ Keldi" if last["action"] == "in" else "🚪 Ketdi"
        status = f"\n📌 Oxirgi: <b>{a_txt}</b> – {last['time'][11:16]}"

    now  = now_tashkent()
    btns = [
        [InlineKeyboardButton("✅ Keldim",  callback_data=f"in_{worker_id}"),
         InlineKeyboardButton("🚪 Ketdim", callback_data=f"out_{worker_id}")],
        [InlineKeyboardButton("⬅️ Orqaga", callback_data="att_menu")],
    ]
    await q.edit_message_text(
        f"👤 <b>{w['name']}</b>\n{'━'*20}\n"
        f"🕐 Hozir: <b>{now.strftime('%H:%M')}</b>{status}\n\nAmalni tanlang:",
        reply_markup=InlineKeyboardMarkup(btns), parse_mode=ParseMode.HTML,
    )


async def record(q, ctx: ContextTypes.DEFAULT_TYPE, worker_id: int, action: str) -> None:
    now      = now_tashkent()
    time_str = now.strftime("%Y-%m-%d %H:%M:%S")
    date_str = now.strftime("%Y-%m-%d")

    late_min = 0
    is_late  = 0
    if action == "in":
        late_min = late_minutes_calc(time_str)
        is_late  = 1 if late_min > 0 else 0

    with db() as conn:
        last = conn.execute(
            "SELECT action FROM attendance WHERE worker_id=? AND date=? ORDER BY id DESC LIMIT 1",
            (worker_id, date_str),
        ).fetchone()

        if last and last["action"] == action:
            a_txt = "keldi" if action == "in" else "ketdi"
            await q.edit_message_text(
                f"⚠️ Siz allaqachon <b>{a_txt}</b> deb belgilangansiz!",
                reply_markup=kb_back("att_menu"), parse_mode=ParseMode.HTML,
            )
            return

        conn.execute(
            "INSERT INTO attendance(worker_id,action,time,date,is_late,late_minutes) VALUES(?,?,?,?,?,?)",
            (worker_id, action, time_str, date_str, is_late, late_min),
        )
        w = conn.execute(
            "SELECT name, telegram_id, username FROM workers WHERE id=?", (worker_id,)
        ).fetchone()

    a_emoji   = "✅" if action == "in" else "🚪"
    a_txt     = "KELDI" if action == "in" else "KETDI"
    late_line = f"\n⚠️ Kechikish: <b>{hm(late_min)}</b>" if is_late else ""

    await q.edit_message_text(
        f"{a_emoji} <b>{a_txt}</b> deb qayd etildi!\n{'━'*20}\n"
        f"👤 <b>{w['name']}</b>\n"
        f"🕐 Vaqt: <b>{now.strftime('%H:%M:%S')}</b>\n"
        f"📅 Sana: <b>{fmt_date_full(date_str)}</b>{late_line}",
        reply_markup=kb_back("att_menu"), parse_mode=ParseMode.HTML,
    )

    uname = f"@{w['username']}" if w["username"] else "—"
    tg_id = str(w["telegram_id"]) if w["telegram_id"] else "—"
    msg   = (
        f"{a_emoji} <b>XODIM {a_txt}</b>\n{'━'*20}\n"
        f"👤 Ism: <b>{w['name']}</b>\n"
        f"🆔 Telegram ID: <code>{tg_id}</code>\n"
        f"📛 Username: {uname}\n"
        f"🕐 Vaqt: <b>{now.strftime('%H:%M:%S')}</b>\n"
        f"📅 Sana: <b>{fmt_date_full(date_str)}</b>"
    )
    if is_late:
        msg += (
            f"\n\n🔴 <b>KECHIKDI!</b>\n"
            f"⏱ {hm(late_min)} kech qoldi\n"
            f"📍 Ish boshlanishi: {WORK_START_H:02d}:{WORK_START_M:02d}"
        )

    for aid in all_admin_ids():
        try:
            await ctx.bot.send_message(aid, msg, parse_mode=ParseMode.HTML)
        except Exception as e:
            log.warning(f"Admin {aid} ga xabar yuborib bolmadi: {e}")


# ═══════════════════════════════════════════════════════════════
#                        HISOBOTLAR
# ═══════════════════════════════════════════════════════════════

def fetch_attendance(start: str, end: str) -> list:
    with db() as conn:
        rows = conn.execute("""
            SELECT w.name, w.telegram_id, w.username,
                   a.action, a.time, a.date, a.is_late, a.late_minutes
            FROM attendance a
            JOIN workers w ON a.worker_id = w.id
            WHERE a.date BETWEEN ? AND ?
            ORDER BY a.date DESC, a.time ASC
        """, (start, end)).fetchall()
    return [dict(r) for r in rows]


async def show_report(q, period: str) -> None:
    now = now_tashkent()
    td  = now.strftime("%Y-%m-%d")

    if period == "today":
        s, e  = td, td
        title = f"📅 Bugungi hisobot – {fmt_date_full(td)}"
    elif period == "week":
        s     = (now - timedelta(days=now.weekday())).strftime("%Y-%m-%d")
        e     = td
        title = f"📆 Haftalik: {fmt_date_full(s)} – {fmt_date_full(e)}"
    else:
        s     = now.strftime("%Y-%m-01")
        e     = td
        title = f"🗓 {MONTHS[now.month-1]} {now.year} – Oylik hisobot"

    data = fetch_attendance(s, e)
    if not data:
        await q.edit_message_text(
            f"<b>{title}</b>\n\n📭 Ma'lumot yo'q.",
            reply_markup=kb_report(), parse_mode=ParseMode.HTML,
        )
        return

    grp: dict = {}
    for r in data:
        n = r["name"]
        if n not in grp:
            grp[n] = {"in": [], "out": [], "late": 0, "late_min": 0}
        if r["action"] == "in":
            grp[n]["in"].append(r["time"][11:16])
            if r["is_late"]:
                grp[n]["late"]    += 1
                grp[n]["late_min"] += r["late_minutes"]
        else:
            grp[n]["out"].append(r["time"][11:16])

    text       = f"<b>{title}</b>\n{'━'*22}\n\n"
    total_late = 0
    for name, d in grp.items():
        late_info = f" ⚠️ {d['late']} marta ({hm(d['late_min'])})" if d["late"] else ""
        text += (
            f"👤 <b>{name}</b>{late_info}\n"
            f"  ✅ Keldi: {', '.join(d['in']) or '—'}\n"
            f"  🚪 Ketdi: {', '.join(d['out']) or '—'}\n\n"
        )
        total_late += d["late"]

    text += f"{'━'*22}\n👥 Jami: <b>{len(grp)}</b> ishchi | ⚠️ Kechikish: <b>{total_late}</b> ta"

    if len(text) > 4000:
        text = text[:3900] + "\n\n<i>...to'liq ma'lumot CSV/Word faylda</i>"

    await q.edit_message_text(text, reply_markup=kb_report(), parse_mode=ParseMode.HTML)


# ═══════════════════════════════════════════════════════════════
#                         CSV EKSPORT
# ═══════════════════════════════════════════════════════════════

async def send_csv(q, ctx: ContextTypes.DEFAULT_TYPE) -> None:
    now  = now_tashkent()
    data = fetch_attendance(now.strftime("%Y-%m-01"), now.strftime("%Y-%m-%d"))
    if not data:
        await q.edit_message_text("📭 Bu oy uchun ma'lumot yo'q.", reply_markup=kb_report())
        return

    buf = io.StringIO()
    wr  = csv.writer(buf, delimiter=";")
    wr.writerow(["Ism Familiya","Telegram ID","Username",
                 "Amal","Vaqt","Sana","Kechikdimi","Kechikish (daqiqa)"])
    for r in data:
        wr.writerow([
            r["name"],
            r["telegram_id"] or "—",
            f"@{r['username']}" if r["username"] else "—",
            "Keldi" if r["action"] == "in" else "Ketdi",
            r["time"][11:19],
            fmt_date_full(r["date"]),
            "Ha" if r["is_late"] else "Yoq",
            r["late_minutes"] or 0,
        ])

    fname = f"hisobot_{now.strftime('%Y_%m')}.csv"
    await ctx.bot.send_document(
        chat_id=q.from_user.id,
        document=io.BytesIO(buf.getvalue().encode("utf-8-sig")),
        filename=fname,
        caption=(f"📎 <b>CSV Hisobot</b>\n📅 {MONTHS[now.month-1]} {now.year}\n📊 {len(data)} yozuv"),
        parse_mode=ParseMode.HTML,
    )
    await q.edit_message_text("✅ CSV yuborildi!", reply_markup=kb_report())


# ═══════════════════════════════════════════════════════════════
#                        WORD EKSPORT
# ═══════════════════════════════════════════════════════════════

def _shd(cell, color: str) -> None:
    tc   = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd  = OxmlElement("w:shd")
    shd.set(qn("w:val"),   "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"),  color)
    tcPr.append(shd)


async def send_word(q, ctx: ContextTypes.DEFAULT_TYPE) -> None:
    now  = now_tashkent()
    data = fetch_attendance(now.strftime("%Y-%m-01"), now.strftime("%Y-%m-%d"))
    if not data:
        await q.edit_message_text("📭 Bu oy uchun ma'lumot yo'q.", reply_markup=kb_report())
        return

    doc = Document()
    sec = doc.sections[0]
    sec.page_width    = Cm(29.7)
    sec.page_height   = Cm(21)
    sec.left_margin   = Cm(2)
    sec.right_margin  = Cm(2)
    sec.top_margin    = Cm(1.5)
    sec.bottom_margin = Cm(1.5)

    t = doc.add_heading("", 0)
    t.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = t.add_run("ISHXONA DAVOMAT HISOBOTI")
    r.font.size = Pt(16); r.font.bold = True
    r.font.color.rgb = RGBColor(0x1a, 0x73, 0xe8)

    sub = doc.add_paragraph()
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sr  = sub.add_run(f"{MONTHS[now.month-1]} {now.year}  |  Toshkent vaqti (UTC+5)")
    sr.font.size = Pt(11); sr.font.color.rgb = RGBColor(0x55, 0x55, 0x55)

    meta = doc.add_paragraph()
    meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
    mr   = meta.add_run(f"Yaratilgan: {now.strftime('%d.%m.%Y %H:%M')}  |  Jami: {len(data)} yozuv")
    mr.font.size = Pt(9)
    doc.add_paragraph()

    hdrs = ["#","Ism Familiya","Telegram ID","Username","Amal","Vaqt","Sana","Kechikish"]
    tbl  = doc.add_table(rows=1, cols=len(hdrs))
    tbl.style = "Table Grid"
    hr = tbl.rows[0]
    for i, h in enumerate(hdrs):
        cell = hr.cells[i]
        cell.text = h
        _shd(cell, "1a73e8")
        p  = cell.paragraphs[0]; p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        rn = p.runs[0]
        rn.font.bold = True; rn.font.size = Pt(9)
        rn.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)

    for idx, row in enumerate(data, 1):
        tr    = tbl.add_row()
        a_txt = "Keldi" if row["action"] == "in" else "Ketdi"
        uname = f"@{row['username']}" if row["username"] else "—"
        late  = f"Ha ({hm(row['late_minutes'])})" if row["is_late"] else "Yoq"
        vals  = [str(idx), row["name"], str(row["telegram_id"] or "—"),
                 uname, a_txt, row["time"][11:19], fmt_date_full(row["date"]), late]
        bg    = "FDEBD0" if row["is_late"] else ("EBF5FB" if idx % 2 == 0 else "FFFFFF")

        for i, val in enumerate(vals):
            cell = tr.cells[i]
            cell.text = val
            _shd(cell, bg)
            p  = cell.paragraphs[0]; p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            rn = p.runs[0] if p.runs else p.add_run(val)
            rn.font.size = Pt(8)
            if row["is_late"] and i == 7:
                rn.font.color.rgb = RGBColor(0xE7, 0x4C, 0x3C)
                rn.font.bold = True

    widths = [Cm(0.7), Cm(3.8), Cm(2.4), Cm(2.8), Cm(1.8), Cm(1.8), Cm(4.8), Cm(2.9)]
    for row in tbl.rows:
        for i, cell in enumerate(row.cells):
            if i < len(widths):
                cell.width = widths[i]

    doc.add_paragraph()
    sp = doc.add_paragraph()
    total_in   = sum(1 for r in data if r["action"] == "in")
    total_out  = sum(1 for r in data if r["action"] == "out")
    total_late = sum(1 for r in data if r["is_late"])
    workers_n  = len(set(r["name"] for r in data))
    sr2 = sp.add_run(
        f"Ishchilar: {workers_n}  |  Kelish: {total_in}  |  Ketish: {total_out}  |  Kechikish: {total_late}"
    )
    sr2.font.size = Pt(9); sr2.font.italic = True

    fp = doc.add_paragraph()
    fp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    fr = fp.add_run("Ishxona Ish Vaqti Tizimi  |  @ishxona_ishvaqtibot")
    fr.font.size = Pt(8); fr.font.color.rgb = RGBColor(0x99, 0x99, 0x99)

    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)
    fname = f"hisobot_{now.strftime('%Y_%m')}.docx"
    await ctx.bot.send_document(
        chat_id=q.from_user.id,
        document=buf,
        filename=fname,
        caption=(f"📄 <b>Word Hisobot</b>\n📅 {MONTHS[now.month-1]} {now.year}\n📊 {len(data)} yozuv"),
        parse_mode=ParseMode.HTML,
    )
    await q.edit_message_text("✅ Word fayl yuborildi!", reply_markup=kb_report())


# ═══════════════════════════════════════════════════════════════
#                    ISHCHI BOSHQARUV
# ═══════════════════════════════════════════════════════════════

async def show_all_workers(q) -> None:
    with db() as conn:
        rows = conn.execute("""
            SELECT w.id, w.name, w.telegram_id, w.username, w.added_date,
                   COUNT(a.id) AS total
            FROM workers w
            LEFT JOIN attendance a ON w.id = a.worker_id
            WHERE w.is_active = 1
            GROUP BY w.id ORDER BY w.name
        """).fetchall()

    if not rows:
        await q.edit_message_text("👥 Ishchilar yo'q.", reply_markup=kb_back("adm_panel"))
        return

    text = f"👥 <b>Barcha Ishchilar ({len(rows)} ta)</b>\n{'━'*22}\n\n"
    for i, w in enumerate(rows, 1):
        uname = f"@{w['username']}" if w["username"] else "—"
        tid   = str(w["telegram_id"]) if w["telegram_id"] else "—"
        text += (
            f"{i}. <b>{w['name']}</b>\n"
            f"   🆔 <code>{tid}</code>  📛 {uname}\n"
            f"   📅 {fmt_dt(w['added_date'])}  |  📊 {w['total']} qayd\n\n"
        )
    await q.edit_message_text(text, reply_markup=kb_back("adm_panel"), parse_mode=ParseMode.HTML)


async def show_del_workers(q) -> None:
    with db() as conn:
        rows = conn.execute(
            "SELECT id, name FROM workers WHERE is_active=1 ORDER BY name"
        ).fetchall()
    if not rows:
        await q.edit_message_text("👥 Ishchilar yo'q.", reply_markup=kb_back("adm_panel"))
        return
    btns = [[InlineKeyboardButton(f"🗑 {w['name']}", callback_data=f"dw_{w['id']}")]
             for w in rows]
    btns.append([InlineKeyboardButton("⬅️ Orqaga", callback_data="adm_panel")])
    await q.edit_message_text(
        "🗑 <b>Ishchi o'chirish</b>\n\nKimni o'chirmoqchisiz?",
        reply_markup=InlineKeyboardMarkup(btns), parse_mode=ParseMode.HTML,
    )


async def delete_worker(q, worker_id: int) -> None:
    with db() as conn:
        w = conn.execute("SELECT name FROM workers WHERE id=?", (worker_id,)).fetchone()
        conn.execute("UPDATE workers SET is_active=0 WHERE id=?", (worker_id,))
    name = w["name"] if w else str(worker_id)
    await q.edit_message_text(
        f"✅ <b>{name}</b> o'chirildi.\n<i>(Ma'lumotlar saqlanib qoldi)</i>",
        reply_markup=kb_back("adm_panel"), parse_mode=ParseMode.HTML,
    )


# ═══════════════════════════════════════════════════════════════
#                    ADMIN BOSHQARUV
# ═══════════════════════════════════════════════════════════════

async def show_all_admins(q) -> None:
    with db() as conn:
        rows = conn.execute("SELECT * FROM admins ORDER BY role, added_date").fetchall()
    emoji_map = {"superadmin": "👑", "admin": "🔑", "manager": "👨"}
    name_map  = {"superadmin": "Developer", "admin": "Admin", "manager": "Manager"}
    text = f"👥 <b>Adminlar Royxati</b>\n{'━'*22}\n\n"
    for a in rows:
        text += (
            f"{emoji_map.get(a['role'],'👤')} <b>{a['name'] or 'Nomsiz'}</b>"
            f" – {name_map.get(a['role'], a['role'])}\n"
            f"   🆔 <code>{a['telegram_id']}</code>  |  {fmt_dt(a['added_date'])}\n\n"
        )
    await q.edit_message_text(text, reply_markup=kb_back("adm_panel"), parse_mode=ParseMode.HTML)


async def show_del_admins(q, role: str) -> None:
    with db() as conn:
        rows = conn.execute(
            "SELECT telegram_id, name FROM admins WHERE role=?", (role,)
        ).fetchall()
    role_txt = "Admin" if role == "admin" else "Manager"
    if not rows:
        await q.edit_message_text(f"👥 {role_txt}lar yo'q.", reply_markup=kb_back("adm_panel"))
        return
    btns = [
        [InlineKeyboardButton(f"🗑 {a['name'] or a['telegram_id']}",
                              callback_data=f"da_{a['telegram_id']}")]
        for a in rows if a["telegram_id"] != SUPER_ADMIN_ID
    ]
    btns.append([InlineKeyboardButton("⬅️ Orqaga", callback_data="adm_panel")])
    await q.edit_message_text(
        f"🗑 <b>{role_txt} o'chirish</b>\n\nKimni o'chirmoqchisiz?",
        reply_markup=InlineKeyboardMarkup(btns), parse_mode=ParseMode.HTML,
    )


async def delete_admin(q, target_id: int) -> None:
    if target_id == SUPER_ADMIN_ID:
        await q.edit_message_text("❌ Developer o'chirib bolmaydi.", reply_markup=kb_back("adm_panel"))
        return
    with db() as conn:
        a = conn.execute("SELECT name FROM admins WHERE telegram_id=?", (target_id,)).fetchone()
        conn.execute("DELETE FROM admins WHERE telegram_id=?", (target_id,))
    name = a["name"] if a else str(target_id)
    await q.edit_message_text(f"✅ <b>{name}</b> o'chirildi.",
                              reply_markup=kb_back("adm_panel"), parse_mode=ParseMode.HTML)


async def show_warn_list(q) -> None:
    with db() as conn:
        rows = conn.execute(
            "SELECT id, name FROM workers WHERE is_active=1 ORDER BY name"
        ).fetchall()
    if not rows:
        await q.edit_message_text("👥 Ishchilar yo'q.", reply_markup=kb_back("adm_panel"))
        return
    btns = [[InlineKeyboardButton(f"⚠️ {w['name']}", callback_data=f"ww_{w['id']}")]
             for w in rows]
    btns.append([InlineKeyboardButton("⬅️ Orqaga", callback_data="adm_panel")])
    await q.edit_message_text(
        "⚠️ <b>Ogohlantirish yuborish</b>\n\nKimga yuborasiz?",
        reply_markup=InlineKeyboardMarkup(btns), parse_mode=ParseMode.HTML,
    )


# ═══════════════════════════════════════════════════════════════
#                     MATN XABAR HANDLER
# ═══════════════════════════════════════════════════════════════

async def on_message(update: Update, ctx: ContextTypes.DEFAULT_TYPE) -> None:
    uid   = update.effective_user.id
    text  = (update.message.text or "").strip()
    state = ctx.user_data.get(STATE)

    try:
        if state == S_ADD_WORKER:
            if len(text) < 3:
                await update.message.reply_text("❌ Ism juda qisqa. Qayta kiriting:")
                return
            now_s = now_tashkent().strftime("%Y-%m-%d %H:%M:%S")
            with db() as conn:
                conn.execute("INSERT INTO workers(name, added_date) VALUES(?,?)", (text, now_s))
            ctx.user_data.clear()
            await update.message.reply_text(
                f"✅ <b>{text}</b> qo'shildi!",
                reply_markup=kb_admin(uid), parse_mode=ParseMode.HTML,
            )

        elif state == S_ADD_ADMIN:
            try:
                new_id = int(text)
            except ValueError:
                await update.message.reply_text("❌ Faqat raqam kiriting:")
                return
            now_s = now_tashkent().strftime("%Y-%m-%d %H:%M:%S")
            with db() as conn:
                conn.execute(
                    "INSERT OR REPLACE INTO admins(telegram_id,name,role,added_date) VALUES(?,?,?,?)",
                    (new_id, f"Admin_{new_id}", "admin", now_s),
                )
            ctx.user_data.clear()
            await update.message.reply_text(
                f"✅ Admin qo'shildi: <code>{new_id}</code>",
                reply_markup=kb_admin(uid), parse_mode=ParseMode.HTML,
            )
            try:
                await ctx.bot.send_message(
                    new_id,
                    "🎉 Siz <b>Ishxona Bot</b> tizimiga <b>admin</b> sifatida qo'shildingiz!\n/start",
                    parse_mode=ParseMode.HTML,
                )
            except Exception:
                pass

        elif state == S_ADD_MANAGER:
            try:
                new_id = int(text)
            except ValueError:
                await update.message.reply_text("❌ Faqat raqam kiriting:")
                return
            now_s = now_tashkent().strftime("%Y-%m-%d %H:%M:%S")
            with db() as conn:
                conn.execute(
                    "INSERT OR REPLACE INTO admins(telegram_id,name,role,added_date) VALUES(?,?,?,?)",
                    (new_id, f"Manager_{new_id}", "manager", now_s),
                )
            ctx.user_data.clear()
            await update.message.reply_text(
                f"✅ Manager qo'shildi: <code>{new_id}</code>",
                reply_markup=kb_admin(uid), parse_mode=ParseMode.HTML,
            )
            try:
                await ctx.bot.send_message(
                    new_id,
                    "🎉 Siz <b>Ishxona Bot</b> tizimiga <b>manager</b> sifatida qo'shildingiz!\n/start",
                    parse_mode=ParseMode.HTML,
                )
            except Exception:
                pass

        elif state == S_BROADCAST:
            with db() as conn:
                workers = conn.execute(
                    "SELECT telegram_id FROM workers WHERE is_active=1 AND telegram_id IS NOT NULL"
                ).fetchall()
            sent = failed = 0
            msg = (f"📢 <b>Korxona Xabari</b>\n{'━'*20}\n{text}\n\n"
                   f"<i>— Ishxona Boshqaruvi</i>")
            for w in workers:
                try:
                    await ctx.bot.send_message(w["telegram_id"], msg, parse_mode=ParseMode.HTML)
                    sent += 1
                    await asyncio.sleep(0.05)
                except Exception:
                    failed += 1
            ctx.user_data.clear()
            await update.message.reply_text(
                f"✅ Xabar yuborildi!\n📤 {sent} ta  |  ❌ {failed} ta xato",
                reply_markup=kb_admin(uid),
            )

        elif state == S_WARN_MSG:
            wid = ctx.user_data.get(WARN_WORKER_ID)
            if not wid:
                ctx.user_data.clear()
                return
            now_s = now_tashkent().strftime("%Y-%m-%d %H:%M:%S")
            with db() as conn:
                w = conn.execute(
                    "SELECT name, telegram_id FROM workers WHERE id=?", (wid,)
                ).fetchone()
                conn.execute(
                    "INSERT INTO warnings(worker_id,message,sent_date,sent_by) VALUES(?,?,?,?)",
                    (wid, text, now_s, uid),
                )
            ctx.user_data.clear()
            if w and w["telegram_id"]:
                try:
                    await ctx.bot.send_message(
                        w["telegram_id"],
                        f"⚠️ <b>OGOHLANTIRISH</b>\n{'━'*20}\n{text}\n\n<i>— Korxona Boshqaruvi</i>",
                        parse_mode=ParseMode.HTML,
                    )
                except Exception:
                    pass
            await update.message.reply_text(
                f"✅ <b>{w['name'] if w else 'Ishchi'}</b>ga ogohlantirish yuborildi.",
                reply_markup=kb_admin(uid), parse_mode=ParseMode.HTML,
            )

        else:
            await update.message.reply_text("Tugmalardan foydalaning:", reply_markup=kb_main(uid))

    except Exception as e:
        log.exception(f"Xabar xatosi: {e}")
        await update.message.reply_text("❌ Xatolik. /start bosing.")


# ═══════════════════════════════════════════════════════════════
#                    KUNLIK TEKSHIRUV
# ═══════════════════════════════════════════════════════════════

async def daily_check(ctx: ContextTypes.DEFAULT_TYPE) -> None:
    today = today_str()
    with db() as conn:
        rows = conn.execute("""
            SELECT w.name,
                   (SELECT a.time FROM attendance a
                    WHERE a.worker_id=w.id AND a.date=? AND a.action='in'
                    ORDER BY a.id LIMIT 1) AS in_time,
                   (SELECT a.late_minutes FROM attendance a
                    WHERE a.worker_id=w.id AND a.date=? AND a.action='in' AND a.is_late=1
                    ORDER BY a.id LIMIT 1) AS late_min
            FROM workers w WHERE w.is_active=1
        """, (today, today)).fetchall()

    late   = [(r["name"], r["late_min"])  for r in rows if r["in_time"] and r["late_min"]]
    absent = [r["name"]                   for r in rows if not r["in_time"]]
    if not late and not absent:
        return

    msg = f"📋 <b>Davomat Holati</b>\n📅 {fmt_date_full(today)}\n{'━'*22}\n\n"
    if late:
        msg += "⚠️ <b>Kechikganlar:</b>\n"
        for name, m in late:
            msg += f"  • {name} – {hm(m)} kech\n"
        msg += "\n"
    if absent:
        msg += "🔴 <b>Kelmadi:</b>\n"
        for name in absent:
            msg += f"  • {name}\n"

    for aid in all_admin_ids():
        try:
            await ctx.bot.send_message(aid, msg, parse_mode=ParseMode.HTML)
        except Exception as e:
            log.warning(f"Kunlik xabar yuborib bolmadi ({aid}): {e}")


# ═══════════════════════════════════════════════════════════════
#                          MAIN
# ═══════════════════════════════════════════════════════════════

def main() -> None:
    init_db()

    app = Application.builder().token(BOT_TOKEN).build()
    app.add_handler(CommandHandler("start", cmd_start))
    app.add_handler(CallbackQueryHandler(on_callback))
    app.add_handler(MessageHandler(filters.TEXT & ~filters.COMMAND, on_message))

    jq = app.job_queue
    if jq:
        import datetime as dt
        run_time = dt.time(hour=9, minute=5, second=0, tzinfo=TASHKENT_TZ)
        jq.run_daily(daily_check, time=run_time)
        log.info("Kunlik tekshiruv jadvallandi: 09:05 Toshkent")

    log.info("Bot ishga tushdi!")
    print("\n" + "="*50)
    print("   ISHXONA ISH VAQTI BOTI ISHGA TUSHDI")
    print("="*50)
    print(f"   Bot     : @ishxona_ishvaqtibot")
    print(f"   Dev ID  : {SUPER_ADMIN_ID}")
    print(f"   Admin ID: {DEFAULT_ADMIN_ID}")
    print(f"   Ish     : {WORK_START_H:02d}:{WORK_START_M:02d} - {WORK_END_H:02d}:{WORK_END_M:02d}")
    print("="*50 + "\n")

    app.run_polling(allowed_updates=Update.ALL_TYPES, drop_pending_updates=True)


if __name__ == "__main__":
    main()
